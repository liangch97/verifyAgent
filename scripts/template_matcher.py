"""Template matcher: detect whether a contract is based on a school template
and report which clauses have been modified.

Usage:
    from scripts.template_matcher import detect_template_match
    result = detect_template_match(contract_text, contract_path=Path("..."))
    # result = {
    #     "matched": True,
    #     "template_name": "中山大学技术开发（委托）合同-中大为乙方.docx",
    #     "similarity": 0.78,
    #     "clauses": [
    #         {"id": "第一条", "title": "...", "status": "unchanged|modified|added|removed",
    #          "template_excerpt": "...", "contract_excerpt": "...", "diff_ratio": 0.95},
    #         ...
    #     ],
    #     "modified_count": 3,
    #     "summary": "...",
    # }

Detection strategy:
1. Cache template texts (loaded once via python-docx, fall back to xml).
2. Compute SequenceMatcher ratio of normalized contract vs each template;
   pick best, mark matched if ratio >= 0.55 OR clause-overlap >= 0.6.
3. Split both texts into clauses by Chinese ordinal headings ("第一条" etc.).
4. For each template clause, find the best-matching contract clause by ratio;
   report status + diff_ratio.

The matcher is tolerant: it accepts varied heading prefixes
(第X条 / 一、 / （一） / 1. / 1.1) and party-name aliases.
"""

from __future__ import annotations

import os
import re
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any

# ---- Heading patterns (tolerant) ----
# Order matters: longer/more-specific first.
_CLAUSE_HEADING_PATTERNS = [
    re.compile(r"^第[一二三四五六七八九十百零〇\d]+条[\s　:：、.\-]*"),
    re.compile(r"^第[一二三四五六七八九十百零〇\d]+章[\s　:：、.\-]*"),
    re.compile(r"^[一二三四五六七八九十]+[、.\s　]+"),
    re.compile(r"^（[一二三四五六七八九十]+）[\s　]*"),
    re.compile(r"^\([一二三四五六七八九十]+\)[\s　]*"),
    re.compile(r"^[０-９0-9]{1,2}[.．、][\s　]*"),
]

_PARTY_ALIASES = {
    # alias -> canonical role (甲 / 乙)
    "委托方": "甲",
    "研究方": "乙",
    "受托方": "乙",
    "合作方": "乙",
    "项目甲方": "甲",
    "项目乙方": "乙",
    "转让方": "甲",
    "受让方": "乙",
    "服务方": "乙",
    "需求方": "甲",
}

# Anchor phrases EXCLUSIVE to SYSU contract templates. They never appear in
# the national 科技部 templates or in industry-supplied custom contracts.
# Generic markers (中山大学 / 示范文本 / 中大) are intentionally NOT included,
# because SYSU often appears as Party B in custom industry contracts and the
# 科技部 templates also use the phrase "示范文本".
_SYSU_ANCHORS = (
    "【科学研究院】",
    "【科研管理部门】",
    "学校合同管理信息系统",
    "合同管理信息系统",
    "中山大学•深圳",
    "中山大学·深圳",
    "中山大学∙深圳",
    "本合同由【科学研究院】",
    "本合同由【科研管理部门】",
)
_SYSU_ANCHOR_MIN = 1

# Heading prefixes that mark the SYSU template's pre-amble "使用说明". Anything
# strictly before the first 第一条 in either contract or template is treated as
# meta-instruction and dropped before clause-level diff to avoid noise.
_PREAMBLE_TERMINATOR = re.compile(r"^第[一二三四五六七八九十百零〇\d]+条")


def _normalize_text(text: str) -> str:
    """Collapse whitespace, strip blank lines, harmonize punctuation."""
    text = text.replace("\u3000", " ").replace("　", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n", text)
    return text.strip()


def _strip_volatile(text: str) -> str:
    """Remove tokens that always differ between template and filled contract,
    so similarity reflects clause structure, not blanks."""
    # blanks (multiple underscores or spaces inside parentheses)
    text = re.sub(r"_{3,}", "____", text)
    text = re.sub(r"（\s*）", "（）", text)
    text = re.sub(r"\(\s*\)", "()", text)
    # dates
    text = re.sub(r"\d{4}年\d{1,2}月\d{1,2}日", "[日期]", text)
    text = re.sub(r"\d{4}[-/]\d{1,2}[-/]\d{1,2}", "[日期]", text)
    # money amounts
    text = re.sub(r"\d[\d,]*\.?\d*\s*(?:元|万元|人民币|RMB)", "[金额]", text)
    return text


def _looks_like_heading(line: str) -> tuple[bool, str | None]:
    s = line.strip()
    if not s or len(s) > 80:
        return False, None
    for pat in _CLAUSE_HEADING_PATTERNS:
        m = pat.match(s)
        if m:
            return True, m.group(0).strip()
    return False, None


def _strip_preamble(text: str) -> str:
    """Drop everything strictly before the first 第N条 heading.

    SYSU template .docx files start with a numbered '使用说明' block (一、二、三、
    四、五、六、七、) that explains how to fill the template. Those headings are
    NOT contract clauses and inflate the diff with bogus removed/added rows.
    Real contracts seldom carry that block. Drop it if and only if a 第N条 line
    is later present, otherwise keep the original text untouched.
    """
    lines = text.splitlines()
    cut = -1
    for i, raw in enumerate(lines):
        if _PREAMBLE_TERMINATOR.match(raw.strip()):
            cut = i
            break
    if cut <= 0:
        return text
    return "\n".join(lines[cut:])


def _split_clauses(text: str) -> list[dict[str, str]]:
    """Split text into clauses keyed by heading. Returns list of
    {"id": "第一条", "title": "项目名称", "body": "..."}."""
    text = _strip_preamble(text)
    clauses: list[dict[str, str]] = []
    current_id = ""
    current_title = ""
    current_body: list[str] = []

    def flush() -> None:
        if current_id or current_body:
            clauses.append({
                "id": current_id,
                "title": current_title,
                "body": "\n".join(current_body).strip(),
            })

    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            current_body.append("")
            continue
        is_h, prefix = _looks_like_heading(line)
        if is_h and prefix:
            # save previous
            flush()
            current_body = []
            current_id = prefix
            # title is the rest of the heading line
            current_title = line.strip()[len(prefix):].strip(" 　:：、.")
        else:
            current_body.append(line)
    flush()

    # drop the very first "preamble" entry if it has no id
    return [c for c in clauses if c.get("id") or len(c.get("body", "")) > 30]


def _count_anchors(text: str) -> int:
    """Number of distinct SYSU template anchor phrases present in `text`."""
    if not text:
        return 0
    return sum(1 for a in _SYSU_ANCHORS if a in text)


def _docx_to_text(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
        doc = Document(str(path))
        parts: list[str] = []
        for para in doc.paragraphs:
            t = para.text or ""
            if t.strip():
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [(c.text or "").strip() for c in row.cells]
                line = " | ".join(c for c in cells if c)
                if line:
                    parts.append(line)
        return "\n".join(parts)
    except Exception:
        # fallback: read xml directly
        import zipfile
        import xml.etree.ElementTree as ET
        try:
            with zipfile.ZipFile(str(path)) as z:
                with z.open("word/document.xml") as f:
                    tree = ET.parse(f)
            ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            texts = [t.text or "" for t in tree.iter(f"{ns}t")]
            return "\n".join(texts)
        except Exception:
            return ""


def resolve_templates_dir() -> Path | None:
    """Resolve the directory containing template .docx files.
    Order: env CONTRACT_TEMPLATES_DIR -> portable_dir/references/templates ->
    /root/contract-review-openclaw-portable/references/templates."""
    env = os.environ.get("CONTRACT_TEMPLATES_DIR")
    if env:
        p = Path(env)
        if p.is_dir():
            return p
    here = Path(__file__).resolve().parent.parent  # portable root
    cand = here / "references" / "templates"
    if cand.is_dir():
        return cand
    fallback = Path("/root/contract-review-openclaw-portable/references/templates")
    if fallback.is_dir():
        return fallback
    return None


_template_cache: dict[str, str] | None = None


def load_templates(templates_dir: Path | None = None) -> dict[str, str]:
    """Load all .docx templates from the given dir. Cached after first call."""
    global _template_cache
    if _template_cache is not None and templates_dir is None:
        return _template_cache
    target = templates_dir or resolve_templates_dir()
    out: dict[str, str] = {}
    if target and target.is_dir():
        for p in sorted(target.glob("*.docx")):
            txt = _docx_to_text(p)
            if txt.strip():
                out[p.name] = txt
    if templates_dir is None:
        _template_cache = out
    return out


def _similarity(a: str, b: str) -> float:
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, a, b, autojunk=False).ratio()


def _diff_ratio(a: str, b: str) -> float:
    """Symmetric similarity ratio of two texts after stripping volatile bits."""
    return _similarity(_strip_volatile(a), _strip_volatile(b))


def _match_clause_pairs(
    template_clauses: list[dict[str, str]],
    contract_clauses: list[dict[str, str]],
) -> list[dict[str, Any]]:
    """For every template clause, find the best contract clause and report status."""
    used: set[int] = set()
    rows: list[dict[str, Any]] = []
    for tc in template_clauses:
        best_idx = -1
        best_ratio = 0.0
        # Prefer same id when present
        same_id_idx = -1
        if tc["id"]:
            for i, cc in enumerate(contract_clauses):
                if i in used:
                    continue
                if cc["id"] == tc["id"]:
                    same_id_idx = i
                    break
        candidates = ([same_id_idx] if same_id_idx >= 0 else []) + [
            i for i in range(len(contract_clauses)) if i not in used and i != same_id_idx
        ]
        for i in candidates:
            cc = contract_clauses[i]
            r = _diff_ratio(tc["body"] or tc["title"], cc["body"] or cc["title"])
            if r > best_ratio:
                best_ratio = r
                best_idx = i
            # short-circuit
            if best_ratio > 0.95:
                break
        if best_idx >= 0 and best_ratio >= 0.30:
            used.add(best_idx)
            cc = contract_clauses[best_idx]
            if best_ratio >= 0.92:
                status = "unchanged"
            elif best_ratio >= 0.55:
                status = "modified"
            else:
                status = "rewritten"
            rows.append({
                "id": tc["id"],
                "title": tc["title"],
                "status": status,
                "diff_ratio": round(best_ratio, 3),
                "template_excerpt": _excerpt(tc["body"] or tc["title"]),
                "contract_excerpt": _excerpt(cc["body"] or cc["title"]),
                "contract_id": cc["id"],
            })
        else:
            rows.append({
                "id": tc["id"],
                "title": tc["title"],
                "status": "removed",
                "diff_ratio": 0.0,
                "template_excerpt": _excerpt(tc["body"] or tc["title"]),
                "contract_excerpt": "",
                "contract_id": "",
            })
    # also report contract clauses that no template clause matched -> "added"
    for i, cc in enumerate(contract_clauses):
        if i in used:
            continue
        body = cc["body"] or cc["title"]
        if len(body) < 20:
            continue
        rows.append({
            "id": cc["id"],
            "title": cc["title"],
            "status": "added",
            "diff_ratio": 0.0,
            "template_excerpt": "",
            "contract_excerpt": _excerpt(body),
            "contract_id": cc["id"],
        })
    return rows


def _excerpt(text: str, max_len: int = 240) -> str:
    t = re.sub(r"\s+", " ", text).strip()
    if len(t) <= max_len:
        return t
    return t[: max_len - 1] + "…"


def detect_template_match(
    contract_text: str,
    *,
    contract_path: Path | None = None,
    templates_dir: Path | None = None,
    match_threshold: float = 0.65,
    anchor_min: int = _SYSU_ANCHOR_MIN,
) -> dict[str, Any]:
    """Detect whether `contract_text` is based on one of the school templates.
    Returns a structured report; `matched=False` when no template matched.

    Match decision combines a *shape* signal (SequenceMatcher similarity
    after stripping volatile tokens) with a *content* signal (count of
    SYSU template anchor phrases). Both must pass — boilerplate clauses
    alone (甲方/乙方/知识产权/保密/违约/争议解决) are not enough.
    """
    contract_text = _normalize_text(contract_text or "")
    templates = load_templates(templates_dir)
    if not templates:
        return {
            "matched": False,
            "reason": "no_templates_loaded",
            "templates_dir": str(resolve_templates_dir() or ""),
        }
    if not contract_text:
        return {"matched": False, "reason": "empty_contract"}

    contract_norm = _strip_volatile(contract_text)
    best_name = ""
    best_ratio = 0.0
    best_template_text = ""
    for name, tt in templates.items():
        tt_norm = _strip_volatile(_normalize_text(tt))
        # Use a fast clause-overlap proxy on first 4k chars to rank, then full ratio on top-k
        r = _similarity(contract_norm[:8000], tt_norm[:8000])
        if r > best_ratio:
            best_ratio = r
            best_name = name
            best_template_text = tt

    anchor_hits = _count_anchors(contract_text)
    shape_pass = best_ratio >= match_threshold
    anchor_pass = anchor_hits >= anchor_min
    matched = shape_pass and anchor_pass

    result: dict[str, Any] = {
        "matched": matched,
        "template_name": best_name,
        "similarity": round(best_ratio, 3),
        "anchor_hits": anchor_hits,
        "anchor_min": anchor_min,
        "match_threshold": match_threshold,
        "templates_considered": list(templates.keys()),
    }
    if not matched:
        if not shape_pass and not anchor_pass:
            reason = (
                f"non_template_contract: similarity_{best_ratio:.2f}<{match_threshold} "
                f"and anchors_{anchor_hits}<{anchor_min}"
            )
        elif not shape_pass:
            reason = f"best_similarity_{best_ratio:.2f}_below_{match_threshold}"
        else:
            reason = (
                f"insufficient_template_anchors_{anchor_hits}_below_{anchor_min} "
                f"(similarity_{best_ratio:.2f}_ok)"
            )
        result["reason"] = reason
        return result

    # Clause-level diff
    template_clauses = _split_clauses(_normalize_text(best_template_text))
    contract_clauses = _split_clauses(contract_text)
    clauses = _match_clause_pairs(template_clauses, contract_clauses)
    counts = {"unchanged": 0, "modified": 0, "rewritten": 0, "added": 0, "removed": 0}
    for c in clauses:
        s = c.get("status", "")
        if s in counts:
            counts[s] += 1
    result["clauses"] = clauses
    result["counts"] = counts
    result["modified_count"] = counts["modified"] + counts["rewritten"] + counts["added"] + counts["removed"]
    result["summary"] = (
        f"匹配范本：{best_name}（相似度 {best_ratio:.0%}，命中范本特征 {anchor_hits} 项）；"
        f"未改动 {counts['unchanged']} 条，已修改 {counts['modified']} 条，"
        f"重写 {counts['rewritten']} 条，新增 {counts['added']} 条，删除 {counts['removed']} 条。"
    )
    return result


__all__ = ["detect_template_match", "load_templates", "resolve_templates_dir"]
